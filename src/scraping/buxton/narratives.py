from docx import Document
import tempfile
from zipfile import ZipFile
import shutil
from pathlib import Path
from os import mkdir

path = "./narratives/Theme - Chord Kbds.docx"
doc = Document(path)

# IMAGE_EXT = ('png', 'jpeg', 'jpg')
#
# with tempfile.TemporaryDirectory() as working_dir:
#     with ZipFile(path) as working_zip:
#         image_list = [name for name in working_zip.namelist() if any(name.endswith(ext) for ext in IMAGE_EXT)]
#         working_zip.extractall(working_dir, image_list)
#         mkdir("./test")
#         for image in image_list:
#             shutil.copy(Path(working_dir).resolve() / image, "./test")

paragraphs = doc.paragraphs
for i in range(len(paragraphs)):
    print(f"{i}: {paragraphs[i].text}")

# for section in doc.sections:
#     print(section.orientation)

# for shape in doc.inline_shapes:
#     print(shape._inline)

# images = doc.tables[0]
# for row in images.rows:
#     contents = []
#     for cell in row.cells:
#         contents.append(cell.text)
    # print(contents)


